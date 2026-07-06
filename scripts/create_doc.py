import os
import subprocess
import sys

def install_and_run():
    # Install python-docx
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.section import WD_ORIENT
    
    doc = Document()
    
    # Set to landscape to give more room
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    # Reduce margins
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    doc.add_heading('ReconNDS Sample Data', 0)
    
    def add_table(title, headers, rows_data):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        
        # Add rows
        for row_data in rows_data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # 1. User Information
    add_table('1. User Information', 
              ['Username', 'Full Name', 'Email', 'Role', 'Status'],
              [
                  ['superadmin', 'Maria Santos', 'msantos@reconndsprivate.sch', 'Super Admin', 'Active'],
                  ['itadmin01', 'Jose Reyes', 'jreyes@reconndsprivate.sch', 'IT Admin', 'Active'],
                  ['staff.cruz', 'Ana Cruz', 'acruz@reconndsprivate.sch', 'Staff', 'Active']
              ])

    # 2. Network Device Inventory
    add_table('2. Network Device Inventory',
              ['Hostname', 'IP Address', 'MAC Address', 'Vendor', 'Trust Level', 'Open Ports', 'Last Seen'],
              [
                  ['PC-ADMIN-01', '192.168.1.10', 'A4:C3:F0:85:2D:11', 'Dell Inc.', 'Trusted', '80, 443', 'Just now'],
                  ['LAPTOP-FACULTY02', '192.168.1.25', 'B8:27:EB:44:1A:C9', 'HP Inc.', 'Trusted', 'None', '2 min ago'],
                  ['UNKNOWN-DEVICE', '192.168.1.88', 'F2:3D:9A:01:CC:47', 'Unknown', 'Unknown', '22, 8080', '5 min ago'],
                  ['PRINTER-LIB', '192.168.1.55', '00:1A:2B:3C:4D:5E', 'Canon', 'Trusted', '9100', '10 min ago'],
                  ['ROGUE-LAPTOP', '192.168.1.99', 'DE:AD:BE:EF:12:34', 'Unknown', 'Blocked', '21, 23', '1 hr ago']
              ])

    # 3. Packet Capture Logs
    add_table('3. Packet Capture Logs',
              ['Scan Date/Time', 'Subnet Scope', 'Devices Found', 'Active Devices', 'Duration', 'Summary'],
              [
                  ['2026-06-29 08:00:12', '192.168.1.0/24', '5', '4', '3.21s', 'Completed nmap scan on 192.168.1.0/24'],
                  ['2026-06-28 17:45:33', '192.168.1.0/24', '4', '3', '4.58s', 'Completed nmap scan on 192.168.1.0/24'],
                  ['2026-06-28 08:02:10', '192.168.1.0/24', '5', '5', '2.97s', 'Completed nmap scan on 192.168.1.0/24']
              ])

    # 4. Spoofing Detection Logs
    add_table('4. Spoofing Detection Logs',
              ['Timestamp', 'Source IP', 'Spoofed MAC', 'Real MAC', 'Detection Type', 'Severity'],
              [
                  ['2026-06-28 10:14:05', '192.168.1.88', 'F2:3D:9A:01:CC:47', 'A4:C3:F0:85:2D:11', 'ARP Spoofing', 'High'],
                  ['2026-06-27 15:33:22', '192.168.1.99', 'DE:AD:BE:EF:12:34', 'B8:27:EB:44:1A:C9', 'MAC Spoofing', 'Critical'],
                  ['2026-06-26 09:05:47', '192.168.1.77', 'CC:AA:BB:DD:EE:FF', '00:1A:2B:3C:4D:5E', 'IP Spoofing', 'Medium']
              ])

    # 5. Alert Logs
    add_table('5. Alert Logs',
              ['Timestamp', 'Alert Type', 'Affected Device', 'IP Address', 'Severity', 'Status'],
              [
                  ['2026-06-29 08:01:44', 'Unknown Device Detected', 'UNKNOWN-DEVICE', '192.168.1.88', 'High', 'Unresolved'],
                  ['2026-06-28 10:14:06', 'ARP Spoofing Attempt', 'UNKNOWN-DEVICE', '192.168.1.88', 'Critical', 'Resolved'],
                  ['2026-06-27 15:33:23', 'Blocked Device Reconnect', 'ROGUE-LAPTOP', '192.168.1.99', 'High', 'Unresolved']
              ])

    # 6. Network Traffic Statistics
    add_table('6. Network Traffic Statistics',
              ['Timestamp', 'Total Packets', 'Bytes Sent', 'Bytes Received', 'Active Connections', 'Top Protocol'],
              [
                  ['2026-06-29 08:00:00', '24,310', '142.3 MB', '389.7 MB', '4', 'TCP'],
                  ['2026-06-28 17:00:00', '19,855', '98.6 MB', '275.1 MB', '3', 'UDP'],
                  ['2026-06-28 08:00:00', '31,042', '210.4 MB', '502.8 MB', '5', 'TCP']
              ])

    # 7. System Activity Logs
    add_table('7. System Activity Logs',
              ['Timestamp', 'Operator', 'Role', 'Action', 'Target MAC/IP', 'Details'],
              [
                  ['2026-06-29 08:00:12', 'msantos', 'Super Admin', 'Ran subnet scan', '192.168.1.0/24', 'Scheduled morning scan'],
                  ['2026-06-28 10:20:00', 'jreyes', 'IT Admin', 'Changed trust level to Blocked', 'DE:AD:BE:EF:12:34', 'Rogue device isolated'],
                  ['2026-06-28 09:45:11', 'jreyes', 'IT Admin', 'Isolated host', '192.168.1.99', 'Threat containment'],
                  ['2026-06-27 14:10:33', 'acruz', 'Staff', 'Viewed device details', '192.168.1.55', 'Routine check']
              ])

    # 8. Report Information
    add_table('8. Report Information',
              ['Report ID', 'Title', 'Generated By', 'Date', 'Type'],
              [
                  ['RPT-2026-001', 'Weekly Security Summary', 'msantos', '2026-06-29', 'Security Summary'],
                  ['RPT-2026-002', 'Device Inventory Report', 'jreyes', '2026-06-28', 'Device Inventory'],
                  ['RPT-2026-003', 'Threat & Spoofing Analysis', 'msantos', '2026-06-27', 'Threat Analysis']
              ])

    # 9. System Configuration
    add_table('9. System Configuration',
              ['Setting', 'Configured Value'],
              [
                  ['Default Subnet Range', '192.168.1.0/24'],
                  ['Scan Mode', 'On-demand + Scheduled (8:00 AM daily)'],
                  ['Alert Notification Email', 'alerts@reconndsprivate.sch'],
                  ['Session Timeout', '30 minutes'],
                  ['Max Failed Login Attempts', '5'],
                  ['Trusted MAC Auto-Approve', 'Disabled'],
                  ['Log Retention Period', '90 days']
              ])

    # 10. Threat Intelligence Database
    add_table('10. Threat Intelligence Database',
              ['Threat ID', 'Threat Name', 'Type', 'Signature/Indicator', 'Risk Level', 'Recommended Action'],
              [
                  ['THR-001', 'ARP Cache Poisoning', 'ARP Spoofing', 'Duplicate IP-to-MAC mapping detected', 'Critical', 'Block device, alert IT Admin'],
                  ['THR-002', 'MAC Address Flooding', 'MAC Spoofing', 'High volume of unknown MACs in short period', 'High', 'Alert and monitor'],
                  ['THR-003', 'Rogue Device Access', 'Unauthorized Device', 'Unregistered MAC connected to network', 'High', 'Flag, notify, isolate'],
                  ['THR-004', 'Open Telnet Port', 'Vulnerability', 'Port 23 open on unknown device', 'Medium', 'Scan & report'],
                  ['THR-005', 'Repeated Auth Failure', 'Brute Force', '5+ failed logins from same IP', 'Medium', 'Lock account, alert admin']
              ])

    # 11. Dashboard Analytics
    add_table('11. Dashboard Analytics',
              ['Metric', 'Value'],
              [
                  ['Total Discovered Devices', '5'],
                  ['Active Nodes', '4'],
                  ['Trusted Devices', '3'],
                  ['Unknown Devices', '1'],
                  ['Blocked Devices', '1'],
                  ['Open Alerts', '2'],
                  ['Last Subnet Scan', '2026-06-29 08:00:12'],
                  ['Scan Duration', '3.21s'],
                  ['Top Threat This Week', 'ARP Cache Poisoning']
              ])

    doc.save('ReconNDS_Sample_Data.docx')

if __name__ == "__main__":
    install_and_run()
